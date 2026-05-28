import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text, bold_words=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p

def add_bullet(doc, bold_title, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(bold_title)
    run.bold = True
    p.add_run(text)

doc = docx.Document()

# PORTADA
doc.add_heading('TRABAJO TERMINAL II', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FlowCode: Compilador de Diagramas de Flujo a Código C para Dispositivos Móviles\n')
run.bold = True
run.font.size = Pt(16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('\nNúmero del TT: 2026-A038\nGrupo: 8CM1\n\n').font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Entregable 1\nCiclo 4: Motor de Análisis\n')
run3.bold = True
run3.font.size = Pt(16)

doc.add_paragraph('\n\n\n')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('Presenta:\n').bold = True
p4.add_run('GARCÍA QUIROZ GUSTAVO IVAN\n\n')
p4.add_run('Director del TT:\n').bold = True
p4.add_run('Prof. EDMUNDO RENÉ DURÁN CAMARILLO')

doc.add_page_break()

# INTRO
add_title(doc, 'Introducción al Motor de Análisis (Ciclo 4)')
add_paragraph(doc, 'El Motor de Análisis constituye la fase central del conversor FlowCode. En este ciclo se desarrolló el mecanismo encargado de extraer la estructura lógica y sintáctica de los diagramas de flujo proporcionados por el usuario, preparándolos para su posterior traducción a código C. A continuación, se detallan los módulos principales que integran este motor, referenciando los archivos de código correspondientes.')

# 1
add_title(doc, '1. Análisis Léxico (Archivo: lib/compiler/lexical_analyzer.dart)')
add_paragraph(doc, 'Se implementó un analizador encargado de examinar el texto de cada nodo en el diagrama de flujo de forma secuencial. El proceso consiste en agrupar los caracteres en unidades lógicas con significado, conocidas como lexemas o "tokens" (cuya estructura se define en lib/compiler/token.dart).')
add_paragraph(doc, 'Durante esta fase, se reconocen 91 tipos de tokens organizados en categorías fundamentales: palabras reservadas en español equivalentes al lenguaje C (como "entero", "si", "mientras"), identificadores de variables, literales numéricas y de cadena, así como operadores aritméticos, relacionales y lógicos. Este proceso descarta elementos prescindibles, como espacios en blanco, y construye la primera abstracción del contenido ingresado.')

# 2
add_title(doc, '2. Análisis Sintáctico (Archivo: lib/compiler/syntax_analyzer.dart)')
add_paragraph(doc, 'Una vez extraídos los tokens, se desarrolló un analizador sintáctico basado en la técnica de descenso recursivo predictivo. Este módulo verifica que la secuencia de lexemas cumpla estrictamente con las reglas de una gramática formal predefinida.')
add_paragraph(doc, 'Por ejemplo, en un nodo de asignación, la gramática valida la presencia obligatoria de un identificador, seguido de un operador de igualdad y una expresión válida. Como resultado de esta fase, se construye el Árbol de Sintaxis Abstracta (AST, implementado en lib/compiler/ast_nodes.dart), el cual organiza las instrucciones en una jerarquía que refleja claramente la precedencia de operadores y el flujo lógico del programa.')

# 3
add_title(doc, '3. Análisis Semántico (Archivo: lib/compiler/semantic_analyzer.dart)')
add_paragraph(doc, 'Posteriormente, se implementó el análisis semántico con el propósito de asegurar la coherencia lógica de las instrucciones expresadas en el AST. Se verificó la consistencia en el uso de los identificadores, comprobando rigurosamente que toda variable haya sido declarada previamente.')
add_paragraph(doc, 'Asimismo, se integraron verificaciones de compatibilidad de tipos en operaciones y asignaciones (por ejemplo, validando que operaciones aritméticas se apliquen únicamente sobre operandos numéricos). Adicionalmente, se incluyeron mecanismos de advertencia para detectar variables no inicializadas y variables declaradas pero no utilizadas, previniendo comportamientos indefinidos.')

# 4
add_title(doc, '4. Tabla de Símbolos (Archivo: lib/compiler/symbol_table.dart)')
add_paragraph(doc, 'Para dar soporte al análisis semántico, se diseñó e integró la Tabla de Símbolos, actuando como la estructura de datos central del motor de conversión. Este componente almacena los metadatos de todos los identificadores (variables, constantes y parámetros) descubiertos durante el análisis.')
add_paragraph(doc, 'En esta tabla se registra el nombre, el tipo de dato subyacente y la ubicación exacta de la declaración (nodo y línea). La gestión de la tabla permite realizar búsquedas rápidas durante la evaluación de expresiones y sienta las bases para el manejo de ámbitos de alcance local y global en la arquitectura del sistema.')

# 5
add_title(doc, '5. Sistema de Errores (Archivo: lib/compiler/compiler_errors.dart)')
add_paragraph(doc, 'Finalmente, se estructuró un sistema unificado para el manejo de errores. Este módulo captura las anomalías identificadas en cualquiera de las fases previas (léxica, sintáctica o semántica) y genera retroalimentación precisa al usuario.')
add_paragraph(doc, 'El sistema clasifica las incidencias mediante un código y un nivel de severidad (informativo, advertencia, error o fatal). Además de localizar el nodo exacto que originó el fallo, proporciona mensajes descriptivos y sugerencias de corrección en idioma español, facilitando el aprendizaje y la resolución iterativa de los diagramas.')
doc.add_page_break()

# EVIDENCIAS
add_title(doc, 'Evidencias de Integración en la Aplicación')
add_paragraph(doc, 'A continuación, se exponen tres plantillas algorítmicas clásicas visualizadas en la aplicación FlowCode. Para ilustrar a profundidad el funcionamiento del Motor de Análisis, se desglosa el proceso de validación mostrando la interfaz principal junto con los resultados específicos de las fases de análisis para cada diagrama.')

# Factorial Iterativo
add_title(doc, "Plantilla '12. Factorial Iterativo'", level=2)

add_bullet(doc, "Evidencia 1.1 - Diagrama Base: ", "Se ilustra el diagrama del cálculo de un factorial mediante ciclos. En esta captura se espera observar la carga del algoritmo iterativo en el lienzo principal del editor.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA FACTORIAL ITERATIVO ]\n')

add_bullet(doc, "Evidencia 1.2 - Análisis Léxico: ", "Se expone el resultado del escáner léxico. Aquí se aprecia cómo el motor agrupa las entradas del usuario en lexemas, distinguiendo palabras reservadas, números y los operadores matemáticos de acumulación del factorial.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS DE ANÁLISIS LÉXICO (TOKENS) ]\n')

add_bullet(doc, "Evidencia 1.3 - Análisis Sintáctico y Semántico: ", "Se visualizan las comprobaciones de gramática (AST) y coherencia semántica. Se evidencia la validación de las sentencias de control y el estado de la Tabla de Símbolos sin errores para las variables de iteración.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS AST Y TABLA DE SÍMBOLOS ]\n')

# Búsqueda Secuencial
add_title(doc, "Plantilla '14. Búsqueda Secuencial'", level=2)

add_bullet(doc, "Evidencia 2.1 - Diagrama Base: ", "Se expone la estructura visual de la búsqueda secuencial en arreglos, mostrando el manejo de nodos condicionales e iterativos en la aplicación.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA BÚSQUEDA SECUENCIAL ]\n')

add_bullet(doc, "Evidencia 2.2 - Análisis Léxico: ", "Se evidencia el reconocimiento de identificadores para arreglos y variables de índice, comprobando la robustez de la fase de tokenización en la extracción de expresiones lógicas complejas.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS DE ANÁLISIS LÉXICO (TOKENS) ]\n')

add_bullet(doc, "Evidencia 2.3 - Análisis Sintáctico y Semántico: ", "Se muestran los resultados del validador. La imagen constata la verificación gramatical del acceso a arreglos y la validación semántica de tipos en las sentencias lógicas de igualdad.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS AST Y TABLA DE SÍMBOLOS ]\n')

# Ordenamiento Burbuja
add_title(doc, "Plantilla '15. Ordenamiento Burbuja'", level=2)

add_bullet(doc, "Evidencia 3.1 - Diagrama Base: ", "Se presenta el algoritmo de ordenamiento cargado, el cual involucra estructuras de repetición anidadas y variables de intercambio temporal directamente en el lienzo interactivo.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA ORDENAMIENTO BURBUJA ]\n')

add_bullet(doc, "Evidencia 3.2 - Análisis Léxico: ", "Se destaca el aislamiento de tokens para múltiples variables indexadas y reasignaciones, reflejando el correcto funcionamiento del agrupamiento de lexemas en un caso de alta densidad.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS DE ANÁLISIS LÉXICO (TOKENS) ]\n')

add_bullet(doc, "Evidencia 3.3 - Análisis Sintáctico y Semántico: ", "Se documenta cómo el validador semántico y la Tabla de Símbolos procesan las iteraciones anidadas, garantizando que el intercambio de posiciones se evalúe correctamente sin errores de tipado o falta de inicialización.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS AST Y TABLA DE SÍMBOLOS ]')


doc.save(r'C:\Users\ivan-\Documents\GitHub\Trabajo-Terminal\flowdiagramapp\docs\ciclo_4\Reporte_Entregable1_Ciclo4_Detallado.docx')
