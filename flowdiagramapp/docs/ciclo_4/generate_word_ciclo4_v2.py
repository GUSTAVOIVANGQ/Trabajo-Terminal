import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p

doc = docx.Document()

# PORTADA
doc.add_heading('TRABAJO TERMINAL II', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FlowCode: Compilador de Diagramas de Flujo a Código C para Dispositivos Móviles\n')
run.bold = True
run.font.size = Pt(16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('\nNúmero del TT: 2026-A038\nGrupo: 8CM1\n\n').font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Entregable 1\nCiclo 4: Motor de Análisis\n')
run3.bold = True
run3.font.size = Pt(16)

doc.add_paragraph('\n\n\n')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('Presenta:\n').bold = True
p4.add_run('GARCÍA QUIROZ GUSTAVO IVAN\n\n')
p4.add_run('Director del TT:\n').bold = True
p4.add_run('Prof. EDMUNDO RENÉ DURÁN CAMARILLO')

doc.add_page_break()

# INTRO
add_title(doc, 'Introducción al Motor de Análisis (Ciclo 4)')
add_paragraph(doc, 'El Motor de Análisis constituye la fase central del conversor FlowCode. En este ciclo se desarrolló el mecanismo encargado de extraer la estructura lógica y sintáctica de los diagramas de flujo proporcionados por el usuario, preparándolos para su posterior traducción a código C. A continuación, se detallan los módulos principales que integran este motor, referenciando los archivos de código correspondientes.')

# 1
add_title(doc, '1. Análisis Léxico (Archivo: lib/compiler/lexical_analyzer.dart)')
add_paragraph(doc, 'Se implementó un analizador encargado de examinar el texto de cada nodo en el diagrama de flujo de forma secuencial. El proceso consiste en agrupar los caracteres en unidades lógicas con significado, conocidas como lexemas o "tokens" (cuya estructura se define en lib/compiler/token.dart).')
add_paragraph(doc, 'Durante esta fase, se reconocen 91 tipos de tokens organizados en categorías fundamentales: palabras reservadas en español equivalentes al lenguaje C (como "entero", "si", "mientras"), identificadores de variables, literales numéricas y de cadena, así como operadores aritméticos, relacionales y lógicos. Este proceso descarta elementos prescindibles, como espacios en blanco, y construye la primera abstracción del contenido ingresado.')
doc.add_page_break()

# 2
add_title(doc, '2. Análisis Sintáctico (Archivo: lib/compiler/syntax_analyzer.dart)')
add_paragraph(doc, 'Una vez extraídos los tokens, se desarrolló un analizador sintáctico basado en la técnica de descenso recursivo predictivo. Este módulo verifica que la secuencia de lexemas cumpla estrictamente con las reglas de una gramática formal predefinida.')
add_paragraph(doc, 'Por ejemplo, en un nodo de asignación, la gramática valida la presencia obligatoria de un identificador, seguido de un operador de igualdad y una expresión válida. Como resultado de esta fase, se construye el Árbol de Sintaxis Abstracta (AST, implementado en lib/compiler/ast_nodes.dart), el cual organiza las instrucciones en una jerarquía que refleja claramente la precedencia de operadores y el flujo lógico del programa.')
doc.add_page_break()

# 3
add_title(doc, '3. Análisis Semántico (Archivo: lib/compiler/semantic_analyzer.dart)')
add_paragraph(doc, 'Posteriormente, se implementó el análisis semántico con el propósito de asegurar la coherencia lógica de las instrucciones expresadas en el AST. Se verificó la consistencia en el uso de los identificadores, comprobando rigurosamente que toda variable haya sido declarada previamente.')
add_paragraph(doc, 'Asimismo, se integraron verificaciones de compatibilidad de tipos en operaciones y asignaciones (por ejemplo, validando que operaciones aritméticas se apliquen únicamente sobre operandos numéricos). Adicionalmente, se incluyeron mecanismos de advertencia para detectar variables no inicializadas y variables declaradas pero no utilizadas, previniendo comportamientos indefinidos.')
doc.add_page_break()

# 4
add_title(doc, '4. Tabla de Símbolos (Archivo: lib/compiler/symbol_table.dart)')
add_paragraph(doc, 'Para dar soporte al análisis semántico, se diseñó e integró la Tabla de Símbolos, actuando como la estructura de datos central del motor de conversión. Este componente almacena los metadatos de todos los identificadores (variables, constantes y parámetros) descubiertos durante el análisis.')
add_paragraph(doc, 'En esta tabla se registra el nombre, el tipo de dato subyacente y la ubicación exacta de la declaración (nodo y línea). La gestión de la tabla permite realizar búsquedas rápidas durante la evaluación de expresiones y sienta las bases para el manejo de ámbitos de alcance local y global en la arquitectura del sistema.')
doc.add_page_break()

# 5
add_title(doc, '5. Sistema de Errores (Archivo: lib/compiler/compiler_errors.dart)')
add_paragraph(doc, 'Finalmente, se estructuró un sistema unificado para el manejo de errores. Este módulo captura las anomalías identificadas en cualquiera de las fases previas (léxica, sintáctica o semántica) y genera retroalimentación precisa al usuario.')
add_paragraph(doc, 'El sistema clasifica las incidencias mediante un código y un nivel de severidad (informativo, advertencia, error o fatal). Además de localizar el nodo exacto que originó el fallo, proporciona mensajes descriptivos y sugerencias de corrección en idioma español, facilitando el aprendizaje y la resolución iterativa de los diagramas.')
doc.add_page_break()

# EVIDENCIAS
add_title(doc, 'Evidencias de Integración en la Aplicación')
add_paragraph(doc, 'A continuación, se exponen tres plantillas algorítmicas clásicas visualizadas en la aplicación FlowCode. Estas capturas evidencian la capacidad de la interfaz para estructurar diagramas complejos, sobre los cuales el Motor de Análisis ejecuta el ciclo de validación léxica, sintáctica y semántica documentado en los apartados anteriores.')

add_title(doc, "Plantilla '12. Factorial Iterativo'", level=2)
add_paragraph(doc, 'Se ilustra el diagrama del cálculo de un factorial mediante ciclos. En esta captura se espera observar la carga del algoritmo iterativo en el lienzo, reflejando el correcto procesamiento de acumuladores y la validación de las sentencias condicionales de control por parte del analizador.')
add_paragraph(doc, '[ ESPACIO PARA INSERTAR IMAGEN DE FACTORIAL ITERATIVO ]\n')

add_title(doc, "Plantilla '14. Búsqueda Secuencial'", level=2)
add_paragraph(doc, 'Se expone la estructura correspondiente a la búsqueda secuencial en arreglos. Se debe visualizar el manejo de variables de índice y expresiones lógicas de igualdad. La imagen constata la validación sintáctica de arreglos y la compatibilidad semántica de tipos en las decisiones lógicas.')
add_paragraph(doc, '[ ESPACIO PARA INSERTAR IMAGEN DE BÚSQUEDA SECUENCIAL ]\n')

add_title(doc, "Plantilla '15. Ordenamiento Burbuja'", level=2)
add_paragraph(doc, 'Se presenta el algoritmo de ordenamiento que involucra estructuras de repetición anidadas. Se aprecia la representación de variables de intercambio temporal y múltiples reasignaciones indexadas. Esto demuestra cómo la Tabla de Símbolos registra y verifica las iteraciones complejas requeridas por el método de ordenamiento de burbuja.')
add_paragraph(doc, '[ ESPACIO PARA INSERTAR IMAGEN DE ORDENAMIENTO BURBUJA ]')

doc.save(r'C:\Users\ivan-\Documents\GitHub\Trabajo-Terminal\flowdiagramapp\docs\ciclo_4\Reporte_Entregable1_Ciclo4_Formal.docx')
