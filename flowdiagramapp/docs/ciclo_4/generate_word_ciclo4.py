import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text, bold_words=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Simple bold replacer if needed, or just plain text
    p.add_run(text)
    return p

doc = docx.Document()

# PORTADA
doc.add_heading('TRABAJO TERMINAL II', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FlowCode: Compilador de Diagramas de Flujo a Código C para Dispositivos Móviles\n')
run.bold = True
run.font.size = Pt(16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('\nNúmero del TT: 2026-A038\nGrupo: 8CM1\n\n').font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Entregable 1\nCiclo 4: Motor de Análisis\n')
run3.bold = True
run3.font.size = Pt(16)

doc.add_paragraph('\n\n\n')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('Presenta:\n').bold = True
p4.add_run('GARCÍA QUIROZ GUSTAVO IVAN\n\n')
p4.add_run('Director del TT:\n').bold = True
p4.add_run('Prof. EDMUNDO RENÉ DURÁN CAMARILLO')

doc.add_page_break()

# INTRO
add_title(doc, 'Introducción al Motor de Análisis (Ciclo 4)')
add_paragraph(doc, 'El Motor de Análisis es el corazón de FlowCode. Piensa en él como el cerebro capaz de "leer" un diagrama de flujo y entender qué significa cada figura y cada palabra escrita dentro de ellas. Este ciclo se encargó de desarrollar las herramientas necesarias para tomar lo que el usuario dibuja en la pantalla de su celular y convertirlo en información organizada y estructurada, paso fundamental antes de poder traducirlo a código C real. A continuación, explicaremos de forma sencilla los componentes principales que hacen esto posible, asignando cada funcionalidad a su archivo de código fuente correspondiente.')

# 1
add_title(doc, '1. Análisis Léxico: Descifrando las palabras del usuario')
add_paragraph(doc, 'Archivo asociado: lib/compiler/lexical_analyzer.dart y lib/compiler/token.dart')
add_paragraph(doc, 'Imagina que estás leyendo un libro en un idioma nuevo. Antes de entender el significado de una oración completa, primero necesitas identificar cada palabra individual y saber qué tipo de palabra es (un verbo, un sustantivo, un signo de puntuación). El Analizador Léxico hace exactamente esto con el texto que los usuarios ingresan en los bloques de sus diagramas de flujo.')
add_paragraph(doc, 'En lugar de ver el texto como una cadena larga de letras, el archivo lexical_analyzer.dart se encarga de separar el texto carácter por carácter y agruparlos en pequeñas unidades con significado llamadas "tokens" (definidos en token.dart). Por ejemplo, si un usuario escribe "numero = 5" en un bloque de proceso, el analizador léxico no ve solo texto, sino que reconoce tres elementos distintos: un "identificador" (la palabra numero), un "operador de asignación" (el signo igual) y un "número entero" (el 5).')
add_paragraph(doc, 'Además, este archivo es tan inteligente que reconoce comandos en español (como "leer", "mostrar", "entero", "si", "mientras") y sabe que equivalen a instrucciones específicas. Gracias a esto, el sistema ignora los espacios en blanco, avisa si olvidaste cerrar unas comillas en un texto, y deja toda la información limpia y etiquetada para el siguiente paso. Es el equivalente a organizar todas las piezas de un rompecabezas por colores antes de empezar a armarlo.')
doc.add_page_break()

# 2
add_title(doc, '2. Análisis Sintáctico: Construyendo oraciones lógicas')
add_paragraph(doc, 'Archivo asociado: lib/compiler/syntax_analyzer.dart y lib/compiler/ast_nodes.dart')
add_paragraph(doc, 'Si el analizador léxico identifica las palabras, el Analizador Sintáctico verifica que esas palabras formen oraciones lógicas respetando las reglas gramaticales. Por ejemplo, en español, la oración "El corre perro rápido" tiene sentido léxico (son palabras reales), pero no tiene sentido sintáctico porque el orden es incorrecto. En programación, sucede exactamente lo mismo.')
add_paragraph(doc, 'El archivo syntax_analyzer.dart toma las piezas (tokens) obtenidas en el paso anterior y verifica que el orden en que fueron escritas tenga sentido según el tipo de bloque del diagrama. Por ejemplo, si estamos en un bloque de "Decisión" (un rombo), este analizador esperará encontrar una comparación lógica (como "numero > 10") y marcará un error si encuentra algo que no pertenece ahí.')
add_paragraph(doc, 'A medida que verifica las reglas, este código va construyendo algo llamado "Árbol de Sintaxis Abstracta" (cuyos nodos se definen en ast_nodes.dart). Piensa en este árbol como un esquema o mapa mental jerárquico de todo el programa. En lugar de tener líneas de texto planas, ahora tenemos una estructura donde es evidente qué operaciones se realizan primero y cómo se relacionan los datos entre sí. Este árbol es el plano arquitectónico final de lo que el diagrama realmente hace.')
doc.add_page_break()

# 3
add_title(doc, '3. Análisis Semántico: Comprobando el sentido y la coherencia')
add_paragraph(doc, 'Archivo asociado: lib/compiler/semantic_analyzer.dart')
add_paragraph(doc, 'Incluso si una oración tiene palabras reales (Análisis Léxico) y está gramaticalmente bien construida (Análisis Sintáctico), podría no tener sentido en el mundo real. Una oración como "El automóvil bebió agua" es gramaticalmente perfecta, pero lógicamente absurda porque los autos no beben agua. El Analizador Semántico se encarga de atrapar este tipo de absurdos lógicos en el diagrama del usuario.')
add_paragraph(doc, 'En FlowCode, el archivo semantic_analyzer.dart examina el árbol que construimos en el paso anterior y verifica la coherencia lógica de las instrucciones. Algunas de sus tareas más importantes son verificar que no estés intentando sumar una palabra con un número, o asegurándose de que si intentas utilizar una variable llamada "resultado", esta haya sido declarada e inicializada previamente en alguna parte del diagrama.')
add_paragraph(doc, 'Es como un corrector de estilo estricto que evita que el programa falle en la vida real. Si intentas dividir un número entre cero o asignar un texto a una variable numérica, este analizador se dará cuenta y generará las advertencias necesarias. Esto garantiza que el código que eventualmente se generará sea sólido, coherente y funcional.')
doc.add_page_break()

# 4
add_title(doc, '4. Tabla de Símbolos: La memoria a corto y largo plazo')
add_paragraph(doc, 'Archivo asociado: lib/compiler/symbol_table.dart')
add_paragraph(doc, 'Para que todo el motor de análisis funcione de manera inteligente, el sistema necesita recordar qué variables se han creado a lo largo del diagrama, qué tipo de dato almacenan y si ya se les asignó algún valor o no. Esta memoria es gestionada por la Tabla de Símbolos.')
add_paragraph(doc, 'El archivo symbol_table.dart funciona como un registro contable o un directorio telefónico muy detallado. Cada vez que el usuario declara una variable nueva en su diagrama, este archivo abre un "expediente" para ella. Guarda su nombre, si es un número, una letra o un texto, en qué parte del diagrama fue creada y registra si en algún momento se ha usado o no.')
add_paragraph(doc, 'Esto es crucial porque permite que, si en un bloque muy lejano del diagrama el usuario hace referencia a la variable "x", el sistema pueda buscar rápidamente en la Tabla de Símbolos, encontrar su expediente, confirmar que es de tipo entero y permitir la operación. Además, esto ayuda a mantener el orden y detectar variables que fueron creadas pero nunca usadas, ayudando a los estudiantes a mantener sus diagramas limpios y optimizados.')
doc.add_page_break()

# 5
add_title(doc, '5. Sistema de Errores: El guía amigable del usuario')
add_paragraph(doc, 'Archivo asociado: lib/compiler/compiler_errors.dart')
add_paragraph(doc, 'Cuando los estudiantes aprenden a programar, es muy común cometer errores, olvidar declarar variables o escribir mal las expresiones. Para que la aplicación FlowCode sea una verdadera herramienta de aprendizaje, no basta con que el sistema falle; debe explicar claramente qué salió mal y cómo solucionarlo.')
add_paragraph(doc, 'El archivo compiler_errors.dart es el encargado de administrar y catalogar todos los problemas que los analizadores (léxico, sintáctico y semántico) encuentren. No solo lanza un error genérico; en su lugar, clasifica el problema (desde simples advertencias hasta errores fatales), registra en qué nodo exacto del diagrama ocurrió el fallo y genera un mensaje en español comprensible.')
add_paragraph(doc, 'Por ejemplo, si un estudiante olvida cerrar un paréntesis, el sistema de errores recoge esta anomalía detectada por el analizador sintáctico y muestra al usuario un mensaje claro junto con una sugerencia de solución. Este archivo consolida todo el historial de incidentes del proceso de conversión para que el usuario reciba un informe detallado y amigable que le ayude a mejorar su diagrama, haciendo de FlowCode no solo un compilador, sino un tutor virtual.')
doc.add_page_break()

# EVIDENCIAS
add_title(doc, 'Evidencias del Sistema en Aplicación (App Android)')
add_paragraph(doc, 'A continuación, se presentan las capturas de pantalla de la aplicación móvil FlowCode mostrando plantillas algorítmicas clásicas cargadas en el lienzo de trabajo. Estas capturas evidencian que la interfaz gráfica es capaz de visualizar y estructurar adecuadamente la topología de los diagramas, proporcionando la base sobre la cual el Motor de Análisis (ciclo 4) ejecuta sus comprobaciones léxicas, sintácticas y semánticas.')

add_title(doc, "Plantilla '12. Factorial Iterativo'", level=2)
add_paragraph(doc, 'En esta captura se debe observar el diagrama de flujo correspondiente al cálculo de un factorial de forma iterativa cargado en la aplicación. Se aprecian los bloques de inicialización y el ciclo lógico necesario. Esta estructura es la que el Motor de Análisis procesa, identificando las variables de acumulación y control, y comprobando la correcta formación de las expresiones matemáticas en los nodos internos.')
add_paragraph(doc, '[ ESPACIO PARA INSERTAR IMAGEN DE FACTORIAL ITERATIVO ]')

add_title(doc, "Plantilla '14. Búsqueda Secuencial'", level=2)
add_paragraph(doc, 'En esta captura de la plantilla de Búsqueda Secuencial, se debe visualizar la estructura de nodos de decisión e iteración. Esto demuestra cómo se representan en la aplicación los arreglos, las variables índice y el elemento objetivo a buscar. A nivel del Motor de Análisis, esta topología y el contenido de los bloques son tokenizados y analizados para asegurar que las comparaciones y la indexación de arreglos cumplan las reglas del lenguaje subyacente.')
add_paragraph(doc, '[ ESPACIO PARA INSERTAR IMAGEN DE BÚSQUEDA SECUENCIAL ]')

add_title(doc, "Plantilla '15. Ordenamiento Burbuja'", level=2)
add_paragraph(doc, 'Esta captura debe mostrar el algoritmo de Ordenamiento Burbuja cargado en el editor. Aquí se deberían poder apreciar ciclos anidados y variables auxiliares ("temporal"). Para el Motor de Análisis, esta plantilla representa un caso complejo de validación semántica, ya que involucra reasignaciones constantes e intercambio de posiciones en arreglos, elementos que son verificados rigurosamente apoyándose en la Tabla de Símbolos.')
add_paragraph(doc, '[ ESPACIO PARA INSERTAR IMAGEN DE ORDENAMIENTO BURBUJA ]')

doc.save(r'C:\Users\ivan-\Documents\GitHub\Trabajo-Terminal\flowdiagramapp\docs\ciclo_4\Reporte_Entregable1_Ciclo4.docx')
