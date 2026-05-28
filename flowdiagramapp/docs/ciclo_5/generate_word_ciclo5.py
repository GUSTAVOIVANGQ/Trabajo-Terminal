import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p

def add_bullet(doc, bold_title, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(bold_title)
    run.bold = True
    p.add_run(text)

doc = docx.Document()

# PORTADA
doc.add_heading('TRABAJO TERMINAL II', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FlowCode: Compilador de Diagramas de Flujo a Código C para Dispositivos Móviles\n')
run.bold = True
run.font.size = Pt(16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('\nNúmero del TT: 2026-A038\nGrupo: 8CM1\n\n').font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Entregable 2\nCiclo 5: Generador de Código\n')
run3.bold = True
run3.font.size = Pt(16)

doc.add_paragraph('\n\n\n')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('Presenta:\n').bold = True
p4.add_run('GARCÍA QUIROZ GUSTAVO IVAN\n\n')
p4.add_run('Director del TT:\n').bold = True
p4.add_run('Prof. EDMUNDO RENÉ DURÁN CAMARILLO')

doc.add_page_break()

# INTRO
add_title(doc, 'Introducción al Generador de Código (Ciclo 5)')
add_paragraph(doc, 'El Generador de Código constituye la fase final del canal de conversión de FlowCode. En este ciclo se desarrolló el mecanismo encargado de optimizar el Árbol de Sintaxis Abstracta (AST) derivado en etapas previas, para posteriormente emitir un programa en código fuente C completamente funcional y válido. A continuación, se detallan los módulos implementados, referenciando los archivos de código fuente correspondientes.')

# 1
add_title(doc, '1. Representación Intermedia AST (Archivo: lib/compiler/ast_nodes.dart)')
add_paragraph(doc, 'Se consolidó la jerarquía de nodos del Árbol de Sintaxis Abstracta (AST) generada en el ciclo anterior, estableciendo este árbol como la representación intermedia definitiva del sistema. Este componente organiza las expresiones y declaraciones en una estructura jerárquica libre de dependencias superficiales.')
add_paragraph(doc, 'Se implementó el patrón arquitectónico "Visitor" para recorrer de forma genérica los subnodos del árbol. Esta abstracción permite separar la lógica estructural del diagrama de las operaciones posteriores, asegurando que tanto la optimización como la generación de código puedan examinar los elementos sin alterar la definición original de las sentencias.')
doc.add_page_break()

# 2
add_title(doc, '2. Optimización del AST (Archivo: lib/compiler/code_optimizer.dart)')
add_paragraph(doc, 'Se desarrolló un módulo de optimización semánticamente equivalente, encargado de reducir redundancias lógicas en el código del usuario antes de su emisión. El proceso se diseñó con una ejecución multi-pasada configurable hasta en cuatro niveles de agresividad.')
add_paragraph(doc, 'Entre las técnicas implementadas se encuentran el plegado de constantes (resolución matemática anticipada), la eliminación de código inalcanzable (código muerto) y la simplificación algebraica de expresiones. Como resultado, se reducen los tiempos y los ciclos de procesamiento del programa final generado, conservando estrictamente el comportamiento original modelado por el estudiante.')
doc.add_page_break()

# 3
add_title(doc, '3. Generador de Código C (Archivo: lib/compiler/code_generator_advanced.dart)')
add_paragraph(doc, 'Finalmente, se construyó el transpilador responsable de emitir código fuente conforme al estándar C99. El módulo traduce directamente los nodos optimizados y el subconjunto de símbolos ISO 5807 soportados hacia sentencias idiomáticas de C.')
add_paragraph(doc, 'Este componente inserta las cabeceras predeterminadas (stdio.h, stdbool.h), estructura la función "main()", e infiere automáticamente los especificadores de formato requeridos ("%d", "%f") basándose en la Tabla de Símbolos. Adicionalmente, el generador normaliza la escritura de operadores lógicos condicionales y estructura de manera automática los bucles anidados, entregando un código formateado, legible y funcional al usuario.')
doc.add_page_break()

# EVIDENCIAS
add_title(doc, 'Evidencias de Integración en la Aplicación')
add_paragraph(doc, 'A continuación, se exponen tres plantillas algorítmicas clásicas visualizadas en la aplicación FlowCode. Para ilustrar a profundidad el funcionamiento de este ciclo, se desglosa el proceso de traducción, mostrando la estructura base, el efecto del optimizador y la visualización final del código C generado.')

# Factorial Iterativo
add_title(doc, "Plantilla '12. Factorial Iterativo'", level=2)

add_bullet(doc, "Evidencia 1.1 - Diagrama Base: ", "Se ilustra el diagrama del cálculo de un factorial mediante ciclos. En esta captura se espera observar la carga del algoritmo iterativo en el lienzo principal del editor.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA FACTORIAL ITERATIVO ]\n')

add_bullet(doc, "Evidencia 1.2 - Optimización del AST: ", "Se muestra la aplicación de las métricas de optimización sobre las sentencias de iteración y cálculo matemático, reflejando el plegado de constantes donde es aplicable.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS DEL OPTIMIZADOR ]\n')

add_bullet(doc, "Evidencia 1.3 - Código C Generado: ", "Se exhibe el texto fuente en lenguaje C emitido por el compilador, evidenciando el mapeo correcto del bucle y la estructura general del archivo.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CÓDIGO C GENERADO ]\n')

# Búsqueda Secuencial
add_title(doc, "Plantilla '14. Búsqueda Secuencial'", level=2)

add_bullet(doc, "Evidencia 2.1 - Diagrama Base: ", "Se expone la estructura visual de la búsqueda secuencial en arreglos, mostrando el manejo de nodos condicionales e iterativos en la aplicación.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA BÚSQUEDA SECUENCIAL ]\n')

add_bullet(doc, "Evidencia 2.2 - Optimización del AST: ", "Se visualizan los resultados del proceso de optimización, destacando la simplificación de condiciones redundantes en el proceso de búsqueda del valor objetivo.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS DEL OPTIMIZADOR ]\n')

add_bullet(doc, "Evidencia 2.3 - Código C Generado: ", "Se comprueba la correcta traducción de los accesos a arreglos y condiciones lógicas al estándar C99, incluyendo especificaciones de formato automático en las funciones printf/scanf.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CÓDIGO C GENERADO ]\n')

# Ordenamiento Burbuja
add_title(doc, "Plantilla '15. Ordenamiento Burbuja'", level=2)

add_bullet(doc, "Evidencia 3.1 - Diagrama Base: ", "Se presenta el algoritmo de ordenamiento cargado, el cual involucra estructuras de repetición anidadas y variables de intercambio temporal.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA ORDENAMIENTO BURBUJA ]\n')

add_bullet(doc, "Evidencia 3.2 - Optimización del AST: ", "Se documenta la simplificación de las expresiones aritméticas y los índices involucrados en el cálculo iterativo del ordenamiento por burbuja.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADOS DEL OPTIMIZADOR ]\n')

add_bullet(doc, "Evidencia 3.3 - Código C Generado: ", "Se presenta el código C resultante, mostrando cómo se estructuraron automáticamente los bucles anidados y el correcto volcado de la variable temporal en las reasignaciones.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CÓDIGO C GENERADO ]')


doc.save(r'C:\Users\ivan-\Documents\GitHub\Trabajo-Terminal\flowdiagramapp\docs\ciclo_5\Reporte_Entregable2_Ciclo5_Formal.docx')
