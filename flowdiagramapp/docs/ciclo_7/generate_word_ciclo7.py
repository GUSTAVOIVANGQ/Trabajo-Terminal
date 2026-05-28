import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p

def add_bullet(doc, bold_title, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(bold_title)
    run.bold = True
    p.add_run(text)

doc = docx.Document()

# PORTADA
doc.add_heading('TRABAJO TERMINAL II', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FlowCode: Compilador de Diagramas de Flujo a Código C para Dispositivos Móviles\n')
run.bold = True
run.font.size = Pt(16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('\nNúmero del TT: 2026-A038\nGrupo: 8CM1\n\n').font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Entregable 4\nCiclo 7 y Entrega Final: Resultados, Rendimiento y Pruebas\n')
run3.bold = True
run3.font.size = Pt(16)

doc.add_paragraph('\n\n\n')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('Presenta:\n').bold = True
p4.add_run('GARCÍA QUIROZ GUSTAVO IVAN\n\n')
p4.add_run('Director del TT:\n').bold = True
p4.add_run('Prof. EDMUNDO RENÉ DURÁN CAMARILLO')

doc.add_page_break()

# INTRO
add_title(doc, 'Introducción a la Validación y Resultados Finales (Ciclo 7)')
add_paragraph(doc, 'Este ciclo representa la etapa final del proyecto FlowCode. Su propósito central fue validar el sistema completo de forma integral mediante pruebas funcionales, de robustez y de rendimiento. Se evaluó rigurosamente el cumplimiento de los criterios de éxito previamente definidos, se documentaron casos de uso reales compilados satisfactoriamente y se delimitaron formalmente las limitaciones y directrices para el trabajo futuro.')

# 1
add_title(doc, '1. Definición de Criterios de Validación Técnica')
add_paragraph(doc, 'Se establecieron métricas objetivas y comprobables vinculadas a los requisitos funcionales (RF) y no funcionales (RNF) del sistema. Esto abarca criterios de corrección funcional (traducción del 100% de los símbolos ISO 5807), calidad del código (estándar C99, indentación consistente) y robustez ante fallos.')
add_paragraph(doc, 'Estas reglas permitieron dictaminar el éxito del proyecto, asegurando que el compilador sea capaz de manejar sentencias malformadas y operaciones inválidas mediante un diagnóstico amigable, sin interrumpir abruptamente el entorno de desarrollo del usuario móvil.')
doc.add_page_break()

# 2
add_title(doc, '2. Resultados Funcionales y Cobertura de Pruebas')
add_paragraph(doc, 'Se documentó la ejecución y aprobación íntegra de una suite automatizada conformada por múltiples escenarios, asegurando una cobertura integral sobre el Motor de Análisis y el Generador de Código.')
add_paragraph(doc, 'Las pruebas confirmaron que las abstracciones generadas (AST y Tabla de Símbolos) mantienen la integridad de los datos durante la transición entre fases. Asimismo, corroboraron la total congruencia lógica al validar desde un esquema visual hasta la estructura final en texto C.')
doc.add_page_break()

# 3
add_title(doc, '3. Rendimiento, Escalabilidad y Tiempos de Respuesta')
add_paragraph(doc, 'Se efectuaron análisis de estrés y métricas de rendimiento para evaluar la velocidad de conversión del pipeline. El objetivo era asegurar un procesamiento fluido imperceptible para el usuario en dispositivos móviles.')
add_paragraph(doc, 'Los análisis concluyeron tiempos promedio en régimen estable significativamente inferiores a un milisegundo (0.80 ms) para topologías regulares, comprobándose una complejidad temporal de O(n) lineal, cumpliendo con los márgenes de operación establecidos por el protocolo original.')
doc.add_page_break()

# 4
add_title(doc, '4. Ejecución de Casos de Uso Reales')
add_paragraph(doc, 'Para evidenciar la operatividad del compilador en entornos prácticos, se plantearon algoritmos fundamentales del aprendizaje de la programación (como factorial, verificación de números primos, búsqueda y ordenamientos).')
add_paragraph(doc, 'Cada esquema visual fue compilado por FlowCode. Posteriormente, el código fuente resultante fue extraído, compilado en GCC (GNU Compiler Collection) sin advertencias de sintaxis y ejecutado para comprobar resultados precisos. Estas validaciones confirmaron plenamente el objetivo central de convertir topologías visuales a programas operativos estandarizados.')
doc.add_page_break()

# 5
add_title(doc, '5. Limitaciones y Trabajo a Futuro')
add_paragraph(doc, 'Alineado a las fronteras acotadas por el protocolo, se establecieron las limitaciones del conversor. En su iteración actual, el sistema excluye arreglos multidimensionales, estructuras dinámicas complejas (structs y memoria dinámica) y subprocesos personalizados por el usuario.')
add_paragraph(doc, 'Dichos apartados se consolidan como oportunidades de expansión futuras. El trabajo próximo incluye incorporar un compilador embebido para probar directamente el ejecutable dentro de la app móvil y posibilitar la traducción hacia otros lenguajes de alto nivel como Python o Java.')
doc.add_page_break()

# EVIDENCIAS
add_title(doc, 'Evidencias Finales: Validación con Algoritmos Reales')
add_paragraph(doc, 'A continuación, se presentan algoritmos representativos procesados íntegramente por FlowCode. Cada plantilla demuestra el flujo real que seguiría el usuario: desde el diseño visual, la consulta del código C emitido, hasta la validación y correcta ejecución del binario en la terminal del sistema.')

# Factorial Iterativo
add_title(doc, "Plantilla '12. Factorial Iterativo'", level=2)

add_bullet(doc, "Evidencia 1.1 - Diagrama Base: ", "Muestra la implementación del cálculo iterativo de un factorial en el lienzo principal de la app.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA FACTORIAL ITERATIVO ]\n')

add_bullet(doc, "Evidencia 1.2 - Diálogo de Resultados: ", "Evidencia la compilación exitosa sin fallos sintácticos ni semánticos, reflejando el código generado listo para uso.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CÓDIGO C DE FACTORIAL ]\n')

add_bullet(doc, "Evidencia 1.3 - Corrida del Ejecutable: ", "Ejecución del binario final compilado con GCC, solicitando datos al usuario y confirmando el cálculo aritmético.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CONSOLA DE EJECUCIÓN (Ej. Factorial de 5 = 120) ]\n')

# Búsqueda Secuencial
add_title(doc, "Plantilla '14. Búsqueda Secuencial'", level=2)

add_bullet(doc, "Evidencia 2.1 - Diagrama Base: ", "Topología algorítmica para la localización de elementos dentro de arreglos estáticos.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA BÚSQUEDA SECUENCIAL ]\n')

add_bullet(doc, "Evidencia 2.2 - Extracción de Código: ", "Traducción fiel de operaciones de acceso a arreglos, condicionales anidadas en ciclos de paro anticipado y funciones printf.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CÓDIGO C DE BÚSQUEDA ]\n')

add_bullet(doc, "Evidencia 2.3 - Corrida del Ejecutable: ", "Terminal validando el caso de coincidencia exitosa al localizar un objetivo específico en memoria.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CONSOLA DE EJECUCIÓN CON BÚSQUEDA EXITOSA ]\n')

# Ordenamiento Burbuja
add_title(doc, "Plantilla '15. Ordenamiento Burbuja'", level=2)

add_bullet(doc, "Evidencia 3.1 - Diagrama Base: ", "Implementación del método de ordenamiento requiriendo variables temporales, intercambio de posiciones y ciclos profundos.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA ORDENAMIENTO BURBUJA ]\n')

add_bullet(doc, "Evidencia 3.2 - Verificación del Compilador GCC: ", "Terminal demostrando que la sentencia 'gcc burbuja.c' transcurre fluidamente y sin lanzar advertencias de incompatibilidad en estándar C99.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CONSOLA DE COMPILACIÓN GCC ]\n')

add_bullet(doc, "Evidencia 3.3 - Corrida del Ejecutable: ", "Comprobación visual de que el algoritmo estructurado por FlowCode logra el ordenamiento final de la secuencia de datos exitosamente.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CONSOLA DE EJECUCIÓN ORDENAMIENTO BURBUJA ]')


doc.save(r'C:\Users\ivan-\Documents\GitHub\Trabajo-Terminal\flowdiagramapp\docs\ciclo_7\Reporte_Entregable4_Ciclo7_Formal.docx')
