import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p

def add_bullet(doc, bold_title, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(bold_title)
    run.bold = True
    p.add_run(text)

doc = docx.Document()

# PORTADA
doc.add_heading('TRABAJO TERMINAL II', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FlowCode: Compilador de Diagramas de Flujo a Código C para Dispositivos Móviles\n')
run.bold = True
run.font.size = Pt(16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('\nNúmero del TT: 2026-A038\nGrupo: 8CM1\n\n').font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Entregable 3\nCiclo 6: Integración y Pruebas\n')
run3.bold = True
run3.font.size = Pt(16)

doc.add_paragraph('\n\n\n')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('Presenta:\n').bold = True
p4.add_run('GARCÍA QUIROZ GUSTAVO IVAN\n\n')
p4.add_run('Director del TT:\n').bold = True
p4.add_run('Prof. EDMUNDO RENÉ DURÁN CAMARILLO')

doc.add_page_break()

# INTRO
add_title(doc, 'Introducción a la Integración y Pruebas (Ciclo 6)')
add_paragraph(doc, 'Este ciclo materializa la consolidación del compilador FlowCode, uniendo las fases individuales de los ciclos anteriores en un canal de conversión (pipeline) completo y funcional. Además de la integración visual y de procesos, se establecieron metodologías de pruebas rigurosas para validar tanto los módulos aislados como el comportamiento del sistema de extremo a extremo.')

# 1
add_title(doc, '1. Orquestación del Flujo de Conversión (Pipeline)')
add_paragraph(doc, 'Se implementó el componente central que coordina la ejecución secuencial del análisis léxico, sintáctico, semántico, la optimización y la generación final de código C. Este pipeline asegura la propagación coherente de los datos entre fases, garantizando que la Tabla de Símbolos y el AST permanezcan sincronizados desde la interpretación del lienzo visual hasta la emisión del texto final.')
doc.add_page_break()

# 2
add_title(doc, '2. Interfaz de Resultados y Sistema de Diagnóstico')
add_paragraph(doc, 'Se desarrolló el componente visual para presentar al usuario los resultados de la compilación. El diálogo interactivo fue dotado de pestañas diferenciadas para exponer las métricas, las unidades léxicas, el AST y el código C con resaltado de sintaxis.')
add_paragraph(doc, 'Simultáneamente, se consolidó el sistema de visualización de errores, incorporando una codificación por color según la severidad (rojo oscuro para errores fatales, naranja para advertencias, etc.), lo cual facilita la depuración y fomenta el autoaprendizaje del estudiante al diseñar sus algoritmos.')
doc.add_page_break()

# 3
add_title(doc, '3. Pruebas Unitarias del Motor de Conversión')
add_paragraph(doc, 'Se estructuró y ejecutó un conjunto de pruebas automatizadas aisladas para garantizar la fiabilidad técnica de cada módulo interno (lexer, parser, validadores semánticos y optimizador).')
add_paragraph(doc, 'Mediante casos predefinidos (positivos y negativos), se verificó el correcto reconocimiento de sentencias, el balance de expresiones lógicas y la prevención de fallos críticos. Estas pruebas demostraron la robustez del motor frente a entradas malformadas, evitando cierres inesperados de la aplicación.')
doc.add_page_break()

# 4
add_title(doc, '4. Pruebas de Integración de Extremo a Extremo (E2E)')
add_paragraph(doc, 'Adicional a las pruebas unitarias, se desarrollaron escenarios automatizados de tipo Extremo a Extremo (End-to-End). Estas pruebas emulan el flujo del usuario desde la inicialización de nodos en el lienzo gráfico hasta la verificación sintáctica y estructural del programa C resultante en el estándar C99.')
add_paragraph(doc, 'La aprobación de estas pruebas constata que las cinco fases del conversor interactúan sin pérdida de información semántica, cubriendo adecuadamente la conversión de los símbolos requeridos del estándar ISO 5807.')
doc.add_page_break()

# 5
add_title(doc, '5. Pruebas Funcionales y Mapeo de Casos de Uso')
add_paragraph(doc, 'Finalmente, se llevaron a cabo pruebas manuales sobre el dispositivo físico destino (p. ej. Samsung Galaxy A26 5G). Esto permitió validar funcionalidades que dependen de componentes externos y de la interfaz de usuario en tiempo de ejecución.')
add_paragraph(doc, 'Entre estas validaciones se verificó el registro de cuentas de usuario, la persistencia local de los proyectos en carpetas y la sincronización de archivos hacia la nube mediante Firebase. Los resultados exitosos trazaron la cobertura completa de los casos de uso documentados en la especificación del sistema.')
doc.add_page_break()

# EVIDENCIAS
add_title(doc, 'Evidencias de Integración y Pruebas Automatizadas')
add_paragraph(doc, 'A continuación, se ilustran escenarios integrales de validación sobre el compilador FlowCode, contrastando la estructura diseñada en la aplicación con los reportes generados y las consolas de prueba.')

# Factorial Iterativo
add_title(doc, "Plantilla '12. Factorial Iterativo'", level=2)

add_bullet(doc, "Evidencia 1.1 - Diagrama Base: ", "Se ilustra el diagrama del cálculo de un factorial mediante ciclos cargado en la aplicación.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA FACTORIAL ITERATIVO ]\n')

add_bullet(doc, "Evidencia 1.2 - Pipeline Completado (Diálogo de Resultados): ", "Se muestra el diálogo de resultados confirmando la ejecución exitosa de todas las fases de análisis para esta plantilla, con el código C generado listado en su pestaña correspondiente.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIÁLOGO DE RESULTADOS / CÓDIGO C ]\n')

add_bullet(doc, "Evidencia 1.3 - Prueba Automatizada: ", "Se documenta la salida en terminal (logs de Flutter/Dart) de la prueba unitaria o de integración correspondiente a la generación de ciclos matemáticos iterativos de forma exitosa.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CONSOLA DE PRUEBAS APROBADAS ]\n')

# Búsqueda Secuencial
add_title(doc, "Plantilla '14. Búsqueda Secuencial'", level=2)

add_bullet(doc, "Evidencia 2.1 - Diagrama Base: ", "Se expone la estructura visual de la búsqueda secuencial en arreglos.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA BÚSQUEDA SECUENCIAL ]\n')

add_bullet(doc, "Evidencia 2.2 - Diagnóstico y Prevención de Errores: ", "Se exhibe la pantalla del sistema de diagnóstico durante la validación del diagrama, reflejando el correcto funcionamiento de los colores de severidad (informativo, advertencias semánticas o errores).")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIÁLOGO DE ERRORES/ADVERTENCIAS ]\n')

add_bullet(doc, "Evidencia 2.3 - Prueba E2E en Terminal: ", "Se captura el reporte de ejecución automatizada donde se avala la conversión íntegra de arreglos y condiciones lógicas de extremo a extremo sin regresiones en el sistema.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: CONSOLA DE PRUEBAS E2E APROBADAS ]\n')

# Ordenamiento Burbuja
add_title(doc, "Plantilla '15. Ordenamiento Burbuja'", level=2)

add_bullet(doc, "Evidencia 3.1 - Diagrama Base: ", "Se presenta el algoritmo complejo de ordenamiento cargado de variables y procesos anidados.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIAGRAMA ORDENAMIENTO BURBUJA ]\n')

add_bullet(doc, "Evidencia 3.2 - Métricas de Conversión: ", "Se ilustra la pestaña de resumen estadístico generada por el pipeline, donde se cuantifica la cantidad de tokens, nodos del AST y tiempo total del proceso para un diagrama de alta densidad.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: DIÁLOGO DE MÉTRICAS ]\n')

add_bullet(doc, "Evidencia 3.3 - Cobertura de Suite Completa: ", "Se presenta la pantalla final del entorno de pruebas evidenciando múltiples casos aprobados, sustentando la fiabilidad de la arquitectura completa documentada en este ciclo.")
add_paragraph(doc, '[ ESPACIO PARA IMAGEN: RESULTADO GENERAL DE TEST RUNNER ]')

doc.save(r'C:\Users\ivan-\Documents\GitHub\Trabajo-Terminal\flowdiagramapp\docs\ciclo_6\Reporte_Entregable3_Ciclo6_Formal.docx')
